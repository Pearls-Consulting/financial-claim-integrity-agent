"""Generate 'Claim Integrity Agent - Trial Demo Guide.docx' in the same house
style as the prequalification trial guide (Calibri, navy headings, one shaded
scope callout, screenshots placed right after the paragraph they illustrate).
Screenshot spots are bordered placeholder boxes to be replaced by hand."""
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(sys.argv[1])

BRAND = RGBColor(0x1F, 0x3A, 0x5F)
GRAY = RGBColor(0x6B, 0x72, 0x80)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
for lvl, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)):
    st = doc.styles[lvl]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = BRAND
    st.font.bold = True


def p(text="", bold=False, italic=False, color=None, size=None, style=None, align=None):
    par = doc.add_paragraph(style=style)
    if align is not None:
        par.alignment = align
    run = par.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return par


def rich(parts, style=None):
    """Paragraph from [(text, bold), ...] so a lead-in can be bold inline."""
    par = doc.add_paragraph(style=style)
    for text, bold in parts:
        r = par.add_run(text)
        r.bold = bold
    return par


def bullet(text, bold_prefix=""):
    par = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = par.add_run(bold_prefix)
        r.bold = True
    par.add_run(text)
    return par


def numbered(text, bold_prefix=""):
    par = doc.add_paragraph(style="List Number")
    if bold_prefix:
        r = par.add_run(bold_prefix)
        r.bold = True
    par.add_run(text)
    return par


def h(text, level):
    doc.add_heading(text, level)


def shot(caption):
    """Bordered placeholder box for a screenshot (replace by hand in Word)."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    cell.width = Inches(6.0)
    par = cell.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before = Pt(28)
    par.paragraph_format.space_after = Pt(28)
    run = par.add_run(f"[ Screenshot — {caption} ]")
    run.italic = True
    run.font.color.rgb = GRAY
    run.font.size = Pt(10)
    doc.add_paragraph()


def note_box(lines, title="Note"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "F2F5F9"})
    tcpr.append(shd)
    trpr = table.rows[0]._tr.get_or_add_trPr()
    trpr.append(trpr.makeelement(qn("w:cantSplit"), {}))
    first = cell.paragraphs[0]
    first.paragraph_format.keep_with_next = True
    r = first.add_run(title)
    r.bold = True
    r.font.color.rgb = BRAND
    for line in lines:
        par = cell.add_paragraph(style="List Bullet")
        par.paragraph_format.keep_with_next = True
        if isinstance(line, tuple):
            lead, rest = line
            rr = par.add_run(lead)
            rr.bold = True
            par.add_run(rest)
        else:
            par.add_run(line)
    doc.add_paragraph()


# ---------------------------------------------------------------- title block
title = p("Claim Integrity Agent Trial Guide", bold=True, size=22, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER)
title.paragraph_format.space_before = Pt(24)
p("AI-Assisted Review of Vendor Claims Before Disbursement", color=GRAY, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
p("Prepared by Pearls Consulting", color=GRAY, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

# ---------------------------------------------------------------- welcome
h("Welcome", 1)
p(
    "This guide walks you through the trial environment of the Claim Integrity Agent: an AI-assisted "
    "workspace that reviews a vendor's payment claim (مطالبة) before it is referred to Finance. The agent "
    "reads the claim's documents — the tax invoice, the contract and Bill of Quantities, the acceptance "
    "document, and the vendor-file attachments — verifies the invoice, cross-checks the documents against "
    "each other and against the contract, and writes a recommendation that cites the rule behind every finding. "
    "The reviewer keeps the decision."
)
p("The guide covers:")
bullet("The claims list — where reviews start and resume;")
bullet("Reviewing a claim — the six guided steps, with what the agent checks at each one;")
bullet("What the agent catches — four situations from day-to-day practice;")
bullet("Scope of the trial and what comes next.")

h("Getting started", 2)
bullet("Access link: provided in the email accompanying this guide.")
bullet(
    "Account: one reviewer login is provided; credentials are shared separately. The trial uses a single "
    "role on purpose — the review is role-agnostic and the same screens serve procurement, vendor "
    "management, finance specialists and executives. Role scopes (who runs which step, who approves) can be "
    "defined per request."
)
bullet(
    "Language: the EN | ع switch in the header flips the whole workspace between English and Arabic, "
    "including right-to-left layout. Findings are written in both languages."
)
bullet(
    "Your review position is saved on the server: if you close the tab, the claim reopens at the step "
    "where you left it."
)
shot("Sign-in page")

note_box(
    [
        (
            "Document intake is by direct upload. ",
            "The “Import from Microsoft Dynamics 365” button shows where the ERP connector attaches in the "
            "full product — it is an indicator of the planned integration, not a live link. The “Claims from "
            "Dynamics 365” table on the home page is sample data shaped like the ERP claims screen.",
        ),
        (
            "AI reading takes time. ",
            "Each uploaded document is read by a specialist OCR engine and then structured by the AI — "
            "typically well under a minute per document, longer for long contracts. If the button shows a "
            "spinner, the read is working, not stuck.",
        ),
        (
            "The agent recommends; the reviewer decides. ",
            "Every finding names its source — a step of the procedure, a ZATCA e-invoicing rule, a VAT "
            "article, the procurement law, or the contract itself — and shows the values it compared. The AI "
            "can never overrule a failed check.",
        ),
        (
            "Rules are configurable. ",
            "The checks in the trial reflect the general regulations and the contract. SDB-specific policy "
            "rules can be added on request without changing the product.",
        ),
    ],
    title="Trial scope",
)

# ---------------------------------------------------------------- 1. claims list
h("1. The Claims List", 1)
p(
    "Sign in and you land on the claims list. The “Claims from Dynamics 365” table mirrors the ERP claims "
    "screen (استلام المطالبات): claim number, vendor, project, contract value (base, excluding VAT), claim "
    "amount including VAT, and a review status. Claims arriving from the ERP open read-only — their fields "
    "and documents come from the ERP record — and are reviewed through the same six steps. The “New claim "
    "review” button starts a review from uploaded documents, which is how you will test the agent during the "
    "trial."
)
p(
    "Each claim shows its review position (which step it is at) and, once a recommendation exists, a "
    "coloured pill: Recommend approve (green), Needs human review (amber), or Recommend reject (red)."
)
shot("Claims list with the sample ERP queue and the “New claim review” button")

# ---------------------------------------------------------------- 2. the review
h("2. Reviewing a Claim", 1)
p(
    "Click “New claim review”. The review runs through six steps shown across the top of the page: Tax "
    "invoice, Contract & BoQ, Acceptance & three-way match, Final check, Pre-finance package, and "
    "Recommendation. Steps 1–5 are review gates; each follows the same rhythm — upload the document, confirm "
    "the fields the agent pre-filled, run the gate, read the findings — then “Continue” moves you to the next "
    "step. Step 6 gathers everything into the recommendation."
)
shot("The six-step review header")

h("How to read the results", 2)
p("Every gate presents its findings the same way, so it is worth knowing the layout once:")
bullet(
    "Each finding is a card with a pass / attention / fail marker, its title in English and Arabic, and a "
    "one-line explanation with the actual figures.",
    bold_prefix="Finding cards. ",
)
bullet(
    "Under every finding, the “Source” line names the rule's origin — for example “VAT Implementing "
    "Regulations, art. 53”, “ZATCA e-invoicing — TLV tags 1–5”, “Procurement Law Executive Regulations — "
    "progress payments”, or “contract — BoQ”.",
    bold_prefix="Source line. ",
)
bullet(
    "“Evidence — values compared” lists the numbers the check used. Any value that was read from a document "
    "carries a “Locate in document” button: the built-in reader opens the PDF on the right page with the "
    "value highlighted, so you can verify what the agent read against the original. Computed values have no "
    "button — the agent never claims a source it does not have.",
    bold_prefix="Evidence and provenance. ",
)
bullet(
    "A check that has nothing to compare (for example, no penalty clause in the contract) says so "
    "explicitly rather than passing silently.",
    bold_prefix="Not-applicable checks. ",
)
shot("A gate's findings with the evidence panel and the PDF reader open on a highlighted value")

# ---- step 1
h("a. Tax invoice (Step 1 — Intake & authenticity)", 2)
p(
    "Upload the vendor's tax invoice PDF. The agent reads it and pre-fills the claim form: vendor name, "
    "invoice number, invoice date, base amount, VAT and total. Complete the claim header — vendor account, "
    "PO and project, contract value (base), payment number, claim type (first payment / periodic / final) "
    "and claim date — and confirm the pre-filled values. Then click “Analyze invoice”."
)
p("The intake gate checks that:")
bullet("the claim fields are complete and the invoice number entered matches the invoice document;")
bullet("the claim total equals base + VAT;")
bullet("the invoice carries a QR code, as a tax invoice must;")
bullet(
    "the QR is genuine ZATCA content and matches the invoice: the agent decodes it in the ZATCA format "
    "(seller name, VAT number, timestamp, total, VAT amount) and compares each value with the invoice — a "
    "decorative or fabricated QR, a VAT number that is not a valid 15-digit ZATCA number, or amounts that "
    "differ from the invoice face are caught here;"
)
bullet(
    "where the invoice carries ZATCA's phase-2 security features (invoice hash, digital signature, public "
    "key), the digital signature is verified mathematically — a signature that does not verify is a "
    "tampering indicator and fails the gate. An invoice with a phase-1 QR only is not failed: the phase-2 "
    "obligation reaches vendors in waves, so the agent flags it for confirmation instead."
)
p(
    "What the QR check can and cannot prove: it establishes that the QR is real ZATCA-format content, that "
    "it agrees with the printed invoice, and — when present — that the cryptographic signature is intact. It "
    "cannot confirm that the invoice was actually reported to or cleared by ZATCA; that requires checking "
    "against ZATCA's own records (for example the Fatoora app), and the agent tells the reviewer exactly "
    "when that confirmation is needed. A QR decoded from the invoice is shown in its own panel with every "
    "field next to the value it was compared with."
)
shot("Step 1 — invoice uploaded, form pre-filled, the QR panel showing decoded fields against the invoice")

# ---- step 2
h("b. Contract & BoQ (Step 2 — BoQ / contract match)", 2)
p(
    "Upload the bank's copy of the contract and Bill of Quantities (one file or two). The agent reads the BoQ "
    "lines and the contract header and suggests the contract value, the contract end date and — when the "
    "contract prints them — its penalty clauses; suggestions are proposed, never applied without you. Enter "
    "the disbursement history for this contract: the amount disbursed before this claim (excluding VAT) and "
    "the number of prior payments. In the full product these come from the ERP payment history. Choose the "
    "contract kind — works / project (acceptance by Certificate of Completion) or goods / supply (acceptance "
    "by goods receipt) — and click “Match against contract”."
)
p("The BoQ gate checks that:")
bullet(
    "every invoice line matches a BoQ line — same item code, same unit price. The line-item table shows all "
    "contracted lines side by side with what this invoice bills; lines not billed this period are shown "
    "muted, which is normal for a periodic claim;"
)
bullet("cumulative claims (disbursed before + this claim, excluding VAT) stay within the contract value;")
bullet("the payment number has not been disbursed before, and follows the prior sequence;")
bullet(
    "the claim type agrees with the payment history — a “final” claim that leaves part of the contract "
    "value unclaimed, or a “first payment” after payments were already made, is flagged with the wording to "
    "return to the vendor (change the claim type to periodic);"
)
bullet("a final claim closes out the contract value exactly.")
shot("Step 2 — contract/BoQ uploaded, suggestions offered, the line-item table comparing BoQ and invoice")

# ---- step 3
h("c. Acceptance & three-way match (Step 3)", 2)
p(
    "Acceptance is evidenced by one document per contract kind: the Certificate of Completion (محضر الإنجاز) "
    "for works, the goods receipt or delivery note for goods. Upload it and click “Run three-way match”. "
    "The agent now reconciles the three sides — what the contract and BoQ agreed, what the acceptance "
    "document says was received, and what the invoice bills."
)
p("The three-way gate checks that:")
bullet("the right acceptance document for the contract kind is present;")
bullet("the Certificate of Completion amount matches the claim total;")
bullet("quantities billed do not exceed quantities received, line by line;")
bullet("quantities received stay within the contracted BoQ quantities;")
bullet("the amount claimed stays within the value of the received work at BoQ prices.")
p(
    "When a claim arrives from the ERP with a posted product receipt, that receipt is used as the received "
    "quantities and cross-checked against the acceptance document. Without a receipt the quantity checks say "
    "so rather than pretending to match."
)
shot("Step 3 — acceptance document uploaded and the three-way findings")

# ---- step 4
h("d. Final check (Step 4 — penalties & delay)", 2)
p(
    "Record any penalties imposed on the vendor for this contract (“Add penalty”: reason, amount, date) — in "
    "the trial these are entered by hand; see the note below on the ERP. Click “Run final check”. This gate "
    "exists because the completion certificate is generated from vendor-fed data and, in practice, a "
    "penalised vendor's certificate stating “no delay” can pass every human gate."
)
p("The final-check gate looks at delay and penalties from three independent angles:")
bullet(
    "the Certificate of Completion's answers (delay? stoppage? observations?) are cross-checked against the "
    "penalty record — “no delay, no stoppage, no observations” while a delay penalty is on record is a "
    "contradiction;",
    bold_prefix="Declared vs recorded. ",
)
bullet(
    "the agent compares the contract end date with the acceptance date itself, independent of what anyone "
    "declared — delay with no penalty on record is flagged for assessment; delay while the certificate says "
    "“no delay” is a contradiction;",
    bold_prefix="Delay inferred from the dates. ",
)
bullet(
    "the penalty clauses read from the uploaded contract (rate, per day or per week, what it applies to, "
    "ceiling, clause reference — each with a locate button into the contract page) are the yardstick: where "
    "the rate applies to the contract value, the expected delay penalty is computed and compared with the "
    "recorded total; a recorded total above the contract's ceiling is flagged.",
    bold_prefix="Penalties vs the contract's own clauses. ",
)
p(
    "What the penalty check can and cannot do: it can read penalty clauses from a contract, detect delay in "
    "delivery or completion from the dates, and measure recorded delay penalties against those clauses. It "
    "cannot detect penalties that depend on information outside the documents — quality, performance or "
    "safety penalties, site observations, or milestone delays that the documents do not date — and where a "
    "clause applies to a line item (“the value of the delayed works”) rather than the contract, the clause "
    "is cited and a penalty demanded, but no figure is invented. The check reads the record it is given; "
    "once we know how penalties are recorded in SDB's ERP, the record is pulled automatically and this gate "
    "runs against it without manual entry."
)
shot("Step 4 — penalties on record, the penalty clauses read from the contract, and the final-check findings")

# ---- step 5
h("e. Pre-finance package (Step 5)", 2)
p(
    "Upload the vendor-file documents: award letter, work commencement minutes, commercial registration, "
    "Zakat certificate and GOSI certificate (the contract and BoQ are already covered by step 2). The agent "
    "identifies each document and lifts its identity fields — CR number, VAT number, reference number, issue "
    "and expiry dates — shown as “Documents reviewed” cards, each with a view button that opens the PDF on "
    "the identifier. The run button enables once the seven required documents are covered; click “Run final "
    "checks”."
)
p("The pre-finance gate checks that:")
bullet("all seven required attachments are on the claim before referral to Finance;")
bullet("VAT is either charged on the invoice or an exemption is declared.")
shot("Step 5 — vendor-file documents identified with their extracted fields")

# ---- step 6
h("f. Recommendation (Step 6)", 2)
p(
    "The last step gathers every gate into one view: the recommendation — Recommend approve, Needs human "
    "review, or Recommend reject — with a bilingual rationale written in the register of a rejection letter, "
    "and one row per gate that opens its findings for an audit walk-through. The wording is drafted by the AI "
    "over the deterministic findings; it can explain and prioritise, but never upgrade a failed check. "
    "“Export matching documents” downloads the documents that took part in the matching as one package. The "
    "reviewer takes the decision and records it in the ERP as today."
)
shot("Step 6 — the recommendation with the gate summary rows")

# ---------------------------------------------------------------- 3. what it catches
h("3. What the Agent Catches", 1)
p(
    "The checks above are easiest to appreciate through the situations they were built for. Each of the "
    "following comes from claims-review practice, and each is caught at a named step."
)

h("A non-VAT invoice, or one with a fake QR", 2)
p(
    "A PDF that merely looks like a tax invoice — no QR, a QR that is decoration, a QR whose VAT number is "
    "not a valid ZATCA number, or one whose amounts do not match the printed totals — fails step 1 before "
    "any amount is examined. An invoice whose phase-2 signature has been altered fails outright; an invoice "
    "with a phase-1 QR only is routed to the reviewer with the exact next action (confirm in the Fatoora "
    "app). A human sees a QR and trusts it; the agent reads it."
)

h("Invoice items against the BoQ — including a claim for work already paid", 2)
p(
    "Step 2 matches every invoice line to the BoQ by item code and unit price, so a re-priced item is "
    "highlighted in both documents. Against the disbursement history, the same step catches a payment "
    "number that was disbursed before, a cumulative total that would exceed the contract value, and a "
    "sequence gap in the payment numbers. Step 3 then closes the remaining door: an invoice with correct "
    "prices and correct totals still fails if it bills more units than were received, or claims more than "
    "the received work is worth at BoQ prices."
)

h("Human errors such as a wrongly generated Certificate of Completion", 2)
p(
    "The certificate is compared with the claim and with the contract rather than taken at face value: an "
    "amount that differs from the claim total, a certificate dated after the contract end while declaring "
    "“no delay”, answers that contradict the penalty record, a works certificate attached to a goods "
    "contract (or a delivery note to a works contract), or received quantities that exceed the BoQ are all "
    "raised as findings with the compared values shown — so the error is corrected before the claim moves "
    "on, not discovered in an audit."
)

h("Claim type — first, periodic, or final", 2)
p(
    "Linked to the past paid claims on the contract (the amount disbursed before and the number of prior "
    "payments), the agent determines whether the claim should be a first payment, a periodic claim, or the "
    "final claim, and flags a mismatch with the wording to return to the vendor — for example, a claim "
    "typed “final” that would leave part of the contract value unclaimed is returned as “change the claim "
    "type to periodic”, the same rejection reason used on the claims screen today. In the full product this "
    "history is read from the ERP, so the check runs without any manual entry."
)

# ---------------------------------------------------------------- 4. scope & next
h("4. Scope of the Trial and What Comes Next", 1)
p("The trial reflects one configuration of the product. The following are scoped for the full deployment:")
bullet(
    "The connector pulls the claim header, the payment history and the attachments from the claims form in "
    "Dynamics 365 Finance & Operations, and the posted product receipt for the three-way match — no "
    "re-entry. The “Import from Microsoft Dynamics 365” button marks where it attaches.",
    bold_prefix="ERP integration. ",
)
bullet(
    "The review gates carry no fixed owner today; procurement, vendor management, finance and executive "
    "roles — and who approves what — are attached per your workflow.",
    bold_prefix="Roles and approval routing. ",
)
bullet(
    "Once penalties are read from the ERP record, the final check runs against them automatically; further "
    "penalty types can be added as their data source is identified.",
    bold_prefix="Penalties from the ERP. ",
)
bullet(
    "Confirmation of an invoice against ZATCA's records (Fatoora) can be added as a policy step for "
    "phase-1 invoices.",
    bold_prefix="Online invoice verification. ",
)
bullet(
    "The checks are rules with cited sources; SDB's own claims policy and contract templates can be encoded "
    "as additional rules without engineering changes.",
    bold_prefix="Policy rules. ",
)

h("Questions & Customization", 1)
p(
    "Review gates, rules, approval workflows, branding, integrations (ERP, document systems, email), and "
    "deployment options are all customizable in the full product — contact us to discuss your requirements."
)
p("We look forward to your feedback.", italic=True)

doc.save(OUT)
print("wrote", OUT)
